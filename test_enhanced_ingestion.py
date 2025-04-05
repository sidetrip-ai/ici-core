from pathlib import Path
import docx
from PyPDF2 import PdfWriter, PdfReader
from reportlab.pdfgen import canvas
from io import BytesIO
from ici.utils.file_utils import ingest_directory
from ici.adapters.vector_store import VectorStoreAdapter

def create_test_files():
    """Create test files in different formats."""
    docs_dir = Path("test_data/documents")
    
    # Create a PDF file
    pdf_buffer = BytesIO()
    c = canvas.Canvas(pdf_buffer)
    c.drawString(100, 750, "This is a test PDF document about artificial intelligence.")
    c.drawString(100, 730, "AI systems can process natural language and perform various tasks.")
    c.save()
    
    # Move to the beginning of the BytesIO buffer
    pdf_buffer.seek(0)
    with open(docs_dir / "ai.pdf", "wb") as f:
        f.write(pdf_buffer.getvalue())
    
    # Create a DOCX file
    doc = docx.Document()
    doc.add_paragraph("Data Science and Machine Learning")
    doc.add_paragraph("Data science involves collecting, processing, and analyzing data to extract insights.")
    doc.add_paragraph("Machine learning is a key component that enables predictive modeling.")
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Algorithm"
    table.cell(0, 1).text = "Use Case"
    table.cell(1, 0).text = "Neural Networks"
    table.cell(1, 1).text = "Image Recognition"
    
    doc.save(docs_dir / "data_science.docx")
    
    # Create a text file
    with open(docs_dir / "nlp.txt", "w") as f:
        f.write("""Natural Language Processing (NLP) is a branch of AI.
It helps computers understand and work with human language.
Common NLP tasks include:
- Text classification
- Named Entity Recognition
- Machine Translation""")

def main():
    print("Testing enhanced file ingestion and preprocessing...")
    
    # Create test files
    print("\n1. Creating test files...")
    create_test_files()
    
    # Create vector store
    vector_store = VectorStoreAdapter(use_persistent=True)
    
    # Ingest all files from the documents directory
    print("\n2. Ingesting files from documents directory...")
    document_ids = ingest_directory(
        "test_data/documents",
        file_extensions=[".txt", ".pdf", ".docx"],
        vector_store=vector_store
    )
    print(f"Successfully ingested {len(document_ids)} documents")
    
    # Test different search queries
    test_queries = [
        "What is artificial intelligence used for?",
        "Explain data science and its components",
        "What are some common NLP tasks?",
        "Tell me about neural networks"
    ]
    
    print("\n3. Testing search functionality...")
    for query in test_queries:
        print(f"\nQuery: {query}")
        results = vector_store.search_similar(query, n_results=2)
        
        # Print results
        if results["documents"]:
            for i in range(len(results["documents"])):
                print(f"\nResult {i + 1}:")
                print(f"Document: {results['metadatas'][i]['file_path']}")
                print(f"Content: {results['documents'][i]}")
        else:
            print("No results found.")

if __name__ == "__main__":
    main() 