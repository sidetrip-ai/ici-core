from pathlib import Path
from typing import Dict, List, Optional, Union, Tuple
import logging
from abc import ABC, abstractmethod
import docx
from PyPDF2 import PdfReader
from ..utils.text_processor import TextPreprocessor

logger = logging.getLogger(__name__)

class DocumentMetadata:
    def __init__(self, source: str, file_path: str, file_type: str):
        self.source = source
        self.file_path = file_path
        self.file_type = file_type

class BaseDocumentScraper(ABC):
    """Base class for document scrapers."""
    
    def __init__(self):
        self.text_processor = TextPreprocessor()
    
    @abstractmethod
    def _extract_text(self, file_path: Union[str, Path]) -> str:
        """Extract raw text from document."""
        pass
    
    def scrape(self, file_path: Union[str, Path], preprocess: bool = True) -> tuple[str, DocumentMetadata]:
        """
        Scrape content from a document.
        
        Args:
            file_path: Path to the document to scrape
            preprocess: Whether to preprocess the extracted text
            
        Returns:
            tuple: (content, metadata)
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract raw text
        content = self._extract_text(path)
        
        # Preprocess text if requested
        if preprocess:
            content = self.text_processor.process_text(
                content,
                clean=True,
                remove_stops=False,  # Keep stopwords for better context
                lemmatize=False,     # Keep original words for better readability
                chunk_size=None
            )
        
        metadata = DocumentMetadata(
            source="file",
            file_path=str(path),
            file_type=path.suffix.lower()[1:]  # Remove the dot from extension
        )
        
        return content, metadata

class TextFileScraper(BaseDocumentScraper):
    """Scraper for plain text files."""
    
    def _extract_text(self, file_path: Union[str, Path]) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

class PDFScraper(BaseDocumentScraper):
    """Scraper for PDF files."""
    
    def _extract_text(self, file_path: Union[str, Path]) -> str:
        reader = PdfReader(file_path)
        text = []
        
        for page in reader.pages:
            text.append(page.extract_text())
        
        return '\n'.join(text)

class DocxScraper(BaseDocumentScraper):
    """Scraper for DOCX files."""
    
    def _extract_text(self, file_path: Union[str, Path]) -> str:
        doc = docx.Document(file_path)
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n'.join(text)

class DocumentScraperFactory:
    """Factory for creating document scrapers based on file type."""
    
    _scrapers: Dict[str, BaseDocumentScraper] = {
        ".txt": TextFileScraper(),
        ".pdf": PDFScraper(),
        ".docx": DocxScraper(),
        # Add more scrapers here for different file types
    }
    
    @classmethod
    def get_scraper(cls, file_path: Union[str, Path]) -> BaseDocumentScraper:
        """Get appropriate scraper for the file type."""
        path = Path(file_path)
        file_extension = path.suffix.lower()
        
        scraper = cls._scrapers.get(file_extension)
        if not scraper:
            raise ValueError(f"No scraper available for file type: {file_extension}")
            
        return scraper

def scrape_file(file_path: Union[str, Path], preprocess: bool = True) -> tuple[str, DocumentMetadata]:
    """
    Scrape content from a file using the appropriate scraper.
    
    Args:
        file_path: Path to the file to scrape
        preprocess: Whether to preprocess the extracted text
        
    Returns:
        tuple: (content, metadata)
        
    Raises:
        ValueError: If no scraper is available for the file type
        FileNotFoundError: If the file doesn't exist
    """
    scraper = DocumentScraperFactory.get_scraper(file_path)
    return scraper.scrape(file_path, preprocess=preprocess) 